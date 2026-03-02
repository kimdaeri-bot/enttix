#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""ENTTIX 창업사업화 지원사업 사업계획서 → .docx 생성"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 헬퍼 ──
def font_set(run, size=10, bold=False, color=None):
    run.font.size = Pt(size)
    run.bold = bold
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))

def set_bg(cell, hex_color):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color); tcPr.append(shd)

def set_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for s, v in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        n = OxmlElement(f'w:{s}'); n.set(qn('w:w'), str(v)); n.set(qn('w:type'), 'dxa')
        tcMar.append(n)
    tcPr.append(tcMar)

def write_cell(cell, text, bold=False, align='left', size=10, bg=None, color=None):
    if bg: set_bg(cell, bg)
    p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
    if align == 'center': p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); font_set(r, size, bold, color)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER; set_margins(cell)

def head_c(cell, text, size=10):
    write_cell(cell, text, bold=True, align='center', size=size, bg='E8E8E8')

def sub_c(cell, text, size=9.5):
    write_cell(cell, text, bold=True, align='center', size=size, bg='F5F5F5')

def mv(tbl, col, r0, r1):  # 세로 병합
    cells = [tbl.cell(r, col) for r in range(r0, r1+1)]
    cells[0].merge(cells[-1]); return cells[0]

def mh(tbl, row, c0, c1):  # 가로 병합
    c = tbl.cell(row, c0); c.merge(tbl.cell(row, c1)); return c

def new_tbl(doc, rows, cols, widths=None):
    t = doc.add_table(rows=rows, cols=cols)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for i, w in enumerate(widths):
            if i < len(t.columns): t.columns[i].width = Cm(w)
    return t

def add_section(doc, text):
    p = doc.add_paragraph(); p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('□ ' + text); font_set(r, 11, True)

def add_caption(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); font_set(r, 10, True)

def add_header_row(tbl, headers):
    row = tbl.rows[0]
    for i, h in enumerate(headers): head_c(row.cells[i], h)
    tblHeader = OxmlElement('w:tblHeader')
    row._tr.get_or_add_trPr().append(tblHeader)

def bullet_cell(cell, items, size=10):
    set_margins(cell)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for i, item in enumerate(items):
        if i == 0:
            p = cell.paragraphs[0]
        else:
            p = cell.add_paragraph()
        r = p.add_run('· ' + item); font_set(r, size)

# ═════════════════════════ 문서 생성 ═════════════════════════
doc = Document()
for sec in doc.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0); sec.right_margin = Cm(3.0)
ns = doc.styles['Normal']
ns.font.name = '맑은 고딕'; ns.font.size = Pt(10)
ns._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ── 제목 ──
tp = doc.add_paragraph(); tp.alignment = WD_ALIGN_PARAGRAPH.CENTER; tp.paragraph_format.space_after = Pt(4)
r = tp.add_run('창업사업화 지원사업 사업계획서'); font_set(r, 16, True)
sp = doc.add_paragraph(); sp.alignment = WD_ALIGN_PARAGRAPH.CENTER; sp.paragraph_format.space_after = Pt(16)
r2 = sp.add_run('ENTTIX  —  AI 예약엔진 기반 글로벌 엔터테인먼트 티켓 플랫폼'); font_set(r2, 10, color='555555')

# ══════════════════════════════
# □ 일반현황
# ══════════════════════════════
add_section(doc, '일반현황')
add_caption(doc, '＜ 일반현황 ＞')

t = new_tbl(doc, 2, 2, [3.5, 14.0])
head_c(t.cell(0,0), '창업아이템명')
write_cell(t.cell(0,1), '"해외 공연·스포츠 티켓, ENTTIX가 AI로 직접 예약해드립니다" — AI 예약엔진 기반 글로벌 엔터테인먼트 티켓 플랫폼')
head_c(t.cell(1,0), '산출물\n(협약기간 내)', size=9.5)
bullet_cell(t.cell(1,1), [
    'AI 예약엔진 (자연어 기반 큐레이션·예약 실행 모듈) 1개',
    '글로벌 티켓 API 연동 모듈 (Tixstock·Ticketmaster 등) 1개',
    'ENTTIX 웹 플랫폼 (Next.js 기반) 1개',
    '여행사용 B2B API 공급 모듈 1개',
    '실시간 예약 현황 관리자 대시보드 1개',
    'AI 기반 엔터테인먼트 티켓 예약 관련 특허 출원 1건',
])

add_caption(doc, '＜ 조직구성 현황 (대표자 본인 제외) ＞')
org_data = [
    ('1','팀장','백엔드 및 AI 파이프라인 개발','Node.js·Python 기반 서버 개발 5년 이상, OpenAI·Claude API 연동 경험 3년, RAG 시스템 설계 경력','팀원 (4대보험)'),
    ('2','팀원','프론트엔드 개발','Next.js·React 기반 웹앱 개발 4년, UX 중심 예약 플로우 설계 경험','팀원 (4대보험)'),
    ('3','팀원','글로벌 티켓 API 연동 및 데이터 관리','REST API 연동 3년 이상, Tixstock·Ticketmaster API 실무 경험, 데이터 파이프라인 설계','팀원 (4대보험)'),
    ('4','팀원','AI 모델 설계 및 프롬프트 엔지니어링','LLM 파인튜닝·RAG 설계 경험, 여행·티켓 도메인 AI 서비스 개발 경력','팀원 예정 (\'26.04)'),
    ('5','팀원','글로벌 파트너십 및 B2B 영업','여행사 대상 B2B 세일즈 경험 3년 이상, 항공·크루즈 GSSA 실무 경력','팀원 예정 (\'26.06)'),
    ('6','팀원','마케팅 및 콘텐츠 기획','디지털 마케팅 경험 3년 이상, 해외 엔터테인먼트 콘텐츠 기획','팀원 예정 (\'26.07)'),
]
ot = new_tbl(doc, len(org_data)+1, 5, [1.2, 1.8, 3.2, 7.5, 3.0])
add_header_row(ot, ['순번','직위','담당 업무','보유역량 (경력 및 학력 등)','구성 상태'])
for ri, (n,p2,du,sk,st) in enumerate(org_data):
    r = ri+1
    for c, (tx, al) in enumerate([(n,'center'),(p2,'center'),(du,'left'),(sk,'left'),(st,'center')]):
        write_cell(ot.cell(r,c), tx, align=al)

# ══════════════════════════════
# □ 창업아이템 개요
# ══════════════════════════════
add_section(doc, '창업아이템 개요 (요약)')
add_caption(doc, '＜ 창업아이템 개요 ＞')

ov = new_tbl(doc, 6, 4, [3.0, 5.5, 2.5, 6.5])
head_c(ov.cell(0,0), '명칭'); write_cell(ov.cell(0,1), 'ENTTIX (엔틱스)')
head_c(ov.cell(0,2), '범주'); write_cell(ov.cell(0,3), 'AI 예약엔진 기반 엔터테인먼트 티켓 플랫폼')
overviews = [
    ('아이템 개요','사용자의 텍스트·음성 입력 → AI 에이전트의 조건 분석 및 상품 큐레이션 → 실시간 예약·결제 페이지 즉시 연결 → 발권의 End-to-End 예약 흐름을 구현한 글로벌 엔터테인먼트 티켓 예약 플랫폼'),
    ('핵심 기능','AI 자연어 입력 → 조건 분석 → 상품 큐레이션 → 실시간 예약·결제 → 발권 (End-to-End 단일 흐름 완결)'),
    ('대상 고객','해외여행 계획 중인 개인 여행자 / 공연·스포츠 관람을 원하는 한국인 / 글로벌 티켓을 취급하는 여행사(B2B)'),
    ('고객 혜택','복잡한 해외 티켓 예약 과정 간소화, AI 기반 맞춤 큐레이션, 실시간 재고 확인 및 즉시 발권'),
    ('가격 모델','티켓 판매 수수료 기반 (GMV의 약 10~15%) + B2B API 공급 월정액 구독료 병행'),
]
for i, (label, text) in enumerate(overviews):
    r = i+1
    head_c(ov.cell(r,0), label)
    c = mh(ov, r, 1, 3); write_cell(c, text)

add_caption(doc, '＜ 핵심 기술 및 특징 ＞')
ft_data = [
    ('AI 예약엔진','자연어 기반 조건 분석 및 예약 실행','텍스트·음성 입력 분석 → 원하는 이벤트·좌석 조건 구조화 → 실시간 재고 기반 상품 큐레이션 → 예약 페이지 즉시 연결','AI를 추천 보조가 아닌 예약 실행 주체(Agent)로 설계'),
    ('실시간 데이터\n주입 (RAG)','글로벌 API 실시간 재고·가격 주입','Tixstock 등 글로벌 티켓 API 실시간 데이터를 AI 컨텍스트에 동적 주입 → 실제 예약 가능한 상품만 추천, AI 환각 방지','AI 환각(Hallucination) 방지, 재고 기반 추천'),
    ('예약 매핑\n엔진','AI 추천 결과와 예약 페이지 실시간 연결','AI 큐레이션 결과를 해당 상품의 예약·결제 페이지와 즉시 연결하는 자체 엔진. 검색→추천→예약 간 단절 구조적 제거','추천→예약 이탈 제거, E2E 단일 흐름'),
    ('노출 최적화\n로직','재고·마진율·선호도 복합 반영 랭킹','재고 가용성·마진율·사용자 선호도를 복합 반영한 자체 랭킹 알고리즘으로 전환율과 수익성 동시 최적화','전환율 + 수익성 동시 최적화'),
    ('B2B API\n공급','여행사 대상 AI 예약엔진 화이트레이블','국내외 여행사에 AI 예약엔진을 API 형태로 공급. 여행사가 자사 서비스에 글로벌 티켓 예약 기능 통합 가능','B2B 월정액 구독 수익, 파트너 확장'),
]
ftt = new_tbl(doc, len(ft_data)+1, 4, [2.8, 3.5, 7.5, 3.7])
add_header_row(ftt, ['구분','핵심 기능','설명','특징'])
for ri, (g,f,d,t2) in enumerate(ft_data):
    r = ri+1
    head_c(ftt.cell(r,0), g, size=9.5)
    for c, tx in enumerate([f,d,t2]): write_cell(ftt.cell(r,c+1), tx)

# ══════════════════════════════
# □ 배경 및 필요성
# ══════════════════════════════
add_section(doc, '1-1. 창업아이템 배경 및 필요성')
add_caption(doc, '＜ 문제 인식 (Problem) ＞')

pm = [
    ('시장 문제','통합 플랫폼 부재','해외 스포츠·공연·뮤지컬·전시 등 엔터테인먼트 티켓을 한 곳에서 예약할 수 있는 한국어 통합 플랫폼 사실상 부재'),
    ('시장 문제','높은 탐색 비용','기존 글로벌 OTA 및 티켓 플랫폼은 검색·필터 기반 상품 나열 구조에 머물러 있어 사용자의 탐색 비용 과다 발생 및 구매 전환 단절'),
    ('시장 문제','언어·접근성 장벽','대부분 영문 위주의 해외 플랫폼으로 한국인 여행자의 접근성 및 활용성이 낮음. 사기·미발권 피해 우려로 신뢰 기반 예약 인프라 부재'),
    ('기존 서비스 한계','AI 연동 부재','AI와 결합된 실질적인 \'예약 실행형\' 서비스 사실상 부재. 기존 서비스는 AI를 추천 보조 기능으로만 활용'),
    ('기존 서비스 한계','예약 단절 구조','검색 → 추천 → 예약 과정에서 단절이 발생해 이탈률이 높음. 여러 사이트를 이동하며 최종 예약까지 다수의 단계가 필요'),
    ('기존 서비스 한계','대표자 실증','항공·크루즈 GSSA 운영(아남항공) 경험을 통해 한국 여행자의 해외 엔터테인먼트 티켓 예약 수요 직접 확인. 현재 Tixstock B2B API 계약 체결 및 파일럿 운영 중'),
]
pmt = new_tbl(doc, len(pm)+1, 3, [2.8, 3.5, 11.2])
add_header_row(pmt, ['구분','문제 유형','내용'])
mv(pmt, 0, 1, 3); head_c(pmt.cell(1,0), '시장\n문제', size=9.5)
mv(pmt, 0, 4, 6); head_c(pmt.cell(4,0), '기존 서비스\n한계', size=9.5)
for ri, (g,s,c) in enumerate(pm):
    r = ri+1; sub_c(pmt.cell(r,1), s); write_cell(pmt.cell(r,2), c)

add_caption(doc, '＜ 해결방안 및 개발 목적 ＞')
sol = [
    ('해결 방안','E2E 예약 흐름 구현','텍스트·음성 입력 → AI 조건 분석 → 실시간 재고 기반 큐레이션 → 예약·결제 페이지 즉시 연결 → 발권까지 단절 없는 단일 흐름 구현'),
    ('해결 방안','글로벌 API 연동','Tixstock·Ticketmaster 등 글로벌 티켓 API 연동으로 해외 스포츠·공연·뮤지컬·전시 등 다양한 카테고리 실시간 재고 및 발권 처리'),
    ('해결 방안','RAG 기반 AI 설계','글로벌 API 실시간 데이터를 AI 컨텍스트에 동적 주입 → AI 환각 방지, 실제 예약 가능한 상품만 추천하는 신뢰 기반 서비스'),
    ('해결 방안','개발 준비 현황','Tixstock API 계약 체결 완료, 파일럿 플랫폼(enttix-omega.vercel.app) 운영 중. 기존 항공·크루즈 GSSA 네트워크 활용 초기 B2B 고객 확보 가능'),
    ('개발 목적','예약 접근성 향상','한국인 여행자가 언어 장벽 없이 해외 엔터테인먼트 티켓을 간편하게 예약할 수 있는 한국어 통합 플랫폼 제공'),
    ('개발 목적','구매 전환율 개선','AI 예약엔진을 통해 검색·추천·예약 간 단절을 제거하여 구매 전환율 구조적 개선 및 탐색 비용 최소화'),
    ('개발 목적','여행산업 융합','여행 일정 기반 콘텐츠 소비 연계 → 관광산업과 엔터테인먼트산업을 실질적인 매출 구조로 융합. B2B API 공급으로 여행사 파트너 생태계 구축'),
]
st = new_tbl(doc, len(sol)+1, 3, [2.8, 3.5, 11.2])
add_header_row(st, ['구분','항목','내용'])
mv(st, 0, 1, 4); head_c(st.cell(1,0), '해결\n방안', size=9.5)
mv(st, 0, 5, 7); head_c(st.cell(5,0), '개발\n목적', size=9.5)
for ri, (g,s,c) in enumerate(sol):
    r = ri+1; sub_c(st.cell(r,1), s); write_cell(st.cell(r,2), c)

# ══════════════════════════════
# □ 목표시장 및 사업화 전략
# ══════════════════════════════
add_section(doc, '목표시장 및 사업화 전략')
add_caption(doc, '＜ 목표시장 및 사업화 전략 ＞')

biz = [
    ('시장 현황','목표 시장','국내 해외여행 인구 연 2,000만 명 이상 / 글로벌 엔터테인먼트 티켓 시장 약 700억 달러(2024) / 한국인 해외 공연·스포츠 관람 수요 급증'),
    ('시장 현황','고객 요구사항','· 믿을 수 있는 한국어 해외 티켓 예약 서비스 필요\n· 복잡한 해외 사이트 없이 원하는 티켓을 한 번에 예약하길 원함\n· 여행 일정과 연계된 맞춤형 엔터테인먼트 추천 필요\n· 실시간 재고 확인 및 즉시 발권 기능 필요'),
    ('비즈니스 모델','수익 구조','① B2C 수수료: 티켓 거래액(GMV)의 약 10~15% 수수료 수취\n② B2B 구독료: 여행사 대상 AI 예약엔진 API 공급, 월정액 구독료'),
    ('비즈니스 모델','고객 확보 전략','초기: 아남항공·디플랫코리아 기존 여행사 네트워크 활용 B2B 선점 → 이후 B2C 직판 확장. 해외 공연·스포츠 커뮤니티 타겟 SNS 마케팅'),
    ('성과 목표','2026년 (협약기간)','B2B 파트너 여행사 5개사 확보 / B2C 월 거래액(GMV) 5,000만 원 이상 / 매출 600,000천원 목표'),
    ('성과 목표','2027년','B2B 파트너 여행사 20개사 / B2C 누적 회원 10,000명 / 매출 1,200,000천원 목표'),
    ('성과 목표','2028년~','글로벌 티켓 카테고리 10개 이상 확장 / 매출 2,000,000천원 이상 / 고용 인력 10명 이상'),
    ('핵심 경쟁력','차별성','단순 추천을 넘어 AI가 직접 예약·결제를 실행하는 End-to-End AI 예약엔진 구현. 국내 유일 AI 기반 글로벌 엔터테인먼트 티켓 통합 예약 플랫폼'),
    ('핵심 경쟁력','기술·사업 경쟁력','Tixstock API 계약 완료 및 파일럿 운영 중(선점 우위), 항공·크루즈 GSSA 운영을 통한 여행 유통 실무 경험, 자체 RAG·예약매핑 엔진 개발'),
]
bt = new_tbl(doc, len(biz)+1, 3, [2.8, 3.5, 11.2])
add_header_row(bt, ['구분','항목','내용'])
mv(bt, 0, 1, 2); head_c(bt.cell(1,0), '시장\n현황', size=9.5)
mv(bt, 0, 3, 4); head_c(bt.cell(3,0), '비즈니스\n모델', size=9.5)
mv(bt, 0, 5, 7); head_c(bt.cell(5,0), '성과\n목표', size=9.5)
mv(bt, 0, 8, 9); head_c(bt.cell(8,0), '핵심\n경쟁력', size=9.5)
for ri, (g,s,c) in enumerate(biz):
    r = ri+1; sub_c(bt.cell(r,1), s); write_cell(bt.cell(r,2), c)

add_caption(doc, '＜ 최종 산출물 작성 (전체 사업단계) ＞')
dl = [
    ('단기\n(협약기간\n내)','AI 예약엔진 (자연어 기반 큐레이션·예약 실행 모듈)','1','2026.06'),
    ('단기\n(협약기간\n내)','글로벌 티켓 API 연동 모듈 (Tixstock·Ticketmaster 등)','1','2026.06'),
    ('단기\n(협약기간\n내)','ENTTIX 웹 플랫폼 (B2C, Next.js 기반)','1','2026.09'),
    ('단기\n(협약기간\n내)','실시간 예약 현황 관리자 대시보드','1','2026.09'),
    ('단기\n(협약기간\n내)','여행사용 B2B API 공급 모듈','1','2026.12'),
    ('단기\n(협약기간\n내)','AI 기반 엔터테인먼트 티켓 예약 관련 특허 출원','1','2026.12'),
    ('중·장기\n최종\n산출물','노출 최적화 랭킹 알고리즘 고도화 모듈','1','2027.06'),
    ('중·장기\n최종\n산출물','글로벌 티켓 카테고리 확장 (10개 이상 카테고리)','1','2027.12'),
    ('중·장기\n최종\n산출물','ENTTIX 모바일 앱 (iOS, AOS)','1','2027.12'),
    ('중·장기\n최종\n산출물','AI 여행 일정 연계 엔터테인먼트 추천 서비스','1','2028.06'),
]
dlt = new_tbl(doc, len(dl)+1, 4, [2.8, 10.0, 1.5, 2.5])
add_header_row(dlt, ['구분','최종 산출물','수량','완료 일정'])
mv(dlt, 0, 1, 6);  head_c(dlt.cell(1,0), '단기\n(협약기간\n내)', size=9.5)
mv(dlt, 0, 7, 10); head_c(dlt.cell(7,0), '중·장기\n최종\n산출물', size=9.5)
for ri, (g,s,q,d) in enumerate(dl):
    r = ri+1; write_cell(dlt.cell(r,1), s)
    write_cell(dlt.cell(r,2), q, align='center'); write_cell(dlt.cell(r,3), d, align='center')

# 마지막
doc.add_paragraph()
ep = doc.add_paragraph('— ENTTIX 창업사업화 지원사업 사업계획서 —')
ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ep.runs[0]; font_set(r, 9, color='666666')

out = '/Users/kim/.openclaw/workspace/enttix/ENTTIX_사업계획서.docx'
doc.save(out)
print(f'✅ {out}')
