#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""오토케(AutoKe) 사업계획서 → .docx 생성"""

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── 기본 여백 ──
for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)

# ── 기본 폰트 ──
style = doc.styles['Normal']
style.font.name = '맑은 고딕'
style.font.size = Pt(10)
style._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

# ═══════════════════ 헬퍼 함수 ═══════════════════

def set_cell_bg(cell, hex_color):
    """셀 배경색 설정"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def set_cell_margins(cell, top=60, bottom=60, left=100, right=100):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top',top),('bottom',bottom),('left',left),('right',right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_repeat_header(row):
    """표 헤더 행 반복"""
    tr = row._tr
    trPr = tr.get_or_add_trPr()
    tblHeader = OxmlElement('w:tblHeader')
    trPr.append(tblHeader)

def add_bold_run(para, text, size=10, color=None, italic=False):
    run = para.add_run(text)
    run.bold = True
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    if italic:
        run.italic = True
    return run

def add_run(para, text, size=10, bold=False, color=None):
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    return run

def cell_para(cell, text, bold=False, align='left', size=10, color=None, bullet=False):
    """셀에 단락 추가"""
    if cell.paragraphs and cell.paragraphs[0].text == '':
        para = cell.paragraphs[0]
    else:
        para = cell.add_paragraph()
    if align == 'center':
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    if bullet:
        text = '· ' + text
    run = para.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    run.font.name = '맑은 고딕'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    if color:
        run.font.color.rgb = RGBColor(*bytes.fromhex(color))
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(cell)
    return para

def head_cell(cell, text, size=10):
    """헤더 셀 (회색 배경 + 볼드 중앙정렬)"""
    set_cell_bg(cell, 'E8E8E8')
    cell_para(cell, text, bold=True, align='center', size=size)

def subhead_cell(cell, text, size=9.5):
    """서브헤더 셀 (연회색 + 중앙정렬)"""
    set_cell_bg(cell, 'F5F5F5')
    cell_para(cell, text, bold=True, align='center', size=size)

def merge_v(table, col, row_start, row_end):
    """세로 병합"""
    cells = [table.cell(r, col) for r in range(row_start, row_end+1)]
    cells[0].merge(cells[-1])
    return cells[0]

def set_table_style(table):
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

def section_title(doc, text, level=2):
    """□ 섹션 제목"""
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(12)
    para.paragraph_format.space_after  = Pt(4)
    add_run(para, '□ ' + text, size=11, bold=True)

def caption(doc, text):
    """＜ 표 캡션 ＞"""
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after  = Pt(3)
    add_run(para, text, size=10, bold=True)

# ═══════════════════ 문서 제목 ═══════════════════

title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_para.paragraph_format.space_before = Pt(6)
title_para.paragraph_format.space_after = Pt(4)
add_bold_run(title_para, '창업사업화 지원사업 사업계획서', size=16)

sub_para = doc.add_paragraph()
sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_para.paragraph_format.space_after = Pt(14)
add_run(sub_para, '오토케(AutoKe) — AI 기반 마케팅 자동화 플랫폼', size=10, color='555555')

# ═══════════════════
# □ 일반현황
# ═══════════════════
section_title(doc, '일반현황')

caption(doc, '＜ 일반현황 ＞')
tbl1 = doc.add_table(rows=2, cols=1)
set_table_style(tbl1)

# Row 0: 창업아이템명
r0 = tbl1.rows[0]
r0.cells[0].merge(r0.cells[0])  # already 1 col
# Actually we want 2 cols
tbl1 = doc.add_table(rows=2, cols=2)
set_table_style(tbl1)
# Remove the one above — redo properly
# Actually let me just create from scratch properly

# Re-create — ignore the above partial tables by clearing
# (python-docx doesn't support easy table removal; let me just code properly)

doc2 = Document()
for section in doc2.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3.0)
    section.right_margin  = Cm(3.0)
style2 = doc2.styles['Normal']
style2.font.name = '맑은 고딕'
style2.font.size = Pt(10)
style2._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')

doc = doc2

# ── 제목 ──
tp = doc.add_paragraph()
tp.alignment = WD_ALIGN_PARAGRAPH.CENTER
tp.paragraph_format.space_after = Pt(4)
add_bold_run(tp, '창업사업화 지원사업 사업계획서', size=16)

sp = doc.add_paragraph()
sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
sp.paragraph_format.space_after = Pt(16)
add_run(sp, '오토케(AutoKe)  —  AI 기반 마케팅 자동화 플랫폼', size=10, color='555555')

# ════════════
# □ 일반현황
# ════════════
section_title(doc, '일반현황')
caption(doc, '＜ 일반현황 ＞')

t = doc.add_table(rows=2, cols=2)
set_table_style(t)
t.columns[0].width = Cm(3.5)
t.columns[1].width = Cm(14)

head_cell(t.cell(0,0), '창업아이템명')
cell_para(t.cell(0,1), '"마케팅, 오토케 하지?" — 콘텐츠 생성부터 전략까지 자동화된 초보자용 마케팅 도구')

head_cell(t.cell(1,0), '산출물\n(협약기간 내)', size=9.5)
items_cell = t.cell(1,1)
items = [
    '웹 플랫폼 1개',
    '관리자 대시보드 1개',
    'SNS 연동 기능이 포함된 AI 콘텐츠 생성 모듈 1개',
    '자동 스케줄링 모듈 1개',
    '마케팅 성과 분석 및 AI 전략 추천 모듈 1개',
    'AI 기반 콘텐츠 생성 관련 특허 출원 1건',
]
for i, item in enumerate(items):
    if i == 0:
        cell_para(items_cell, item, bullet=True)
    else:
        p = items_cell.add_paragraph('· ' + item)
        p.paragraph_format.space_before = Pt(0)
        r = p.runs[0]
        r.font.size = Pt(10)
        r.font.name = '맑은 고딕'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), '맑은 고딕')
    items_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    set_cell_margins(items_cell)

# 조직구성 표
caption(doc, '＜ 조직구성 현황 (대표자 본인 제외) ＞')
org_data = [
    ('1','팀장','백엔드 개발','산업디자인 전공, 웹 서버 및 API 개발 5년 이상, AI 서비스 개발 경력 2년','팀원 (4대보험)'),
    ('2','팀원','프론트엔드 개발','UI 설계 및 React 기반 웹앱 개발 4년','팀원 (4대보험)'),
    ('3','팀원','콘텐츠 디자인 및 UI/UX 기획','SNS 콘텐츠 포맷 전문 경험, 스타트업 경력 1년 이상','팀원 (4대보험)'),
    ('4','팀원','콘텐츠 생성 AI 연동 및 모델 적용','GPT · Claude API 연동 경험, AI 서비스 개발','팀원 예정 (\'25.07)'),
    ('5','팀원','커뮤니티 운영 및 SNS 채널 관리','실무 중심 SNS 마케팅 경험 3년 이상','팀원 예정 (\'25.06)'),
    ('6','팀원','고객응대, 이용자 지원','스타트업 운영관리 경험 1년 이상','팀원 예정 (\'25.07)'),
    ('7','팀원','비즈니스 모델 및 사업 기획','창업컨설팅 및 초기 사업 전략 수립 경험','팀원 예정 (\'25.07)'),
]
org = doc.add_table(rows=len(org_data)+1, cols=5)
set_table_style(org)
org.columns[0].width = Cm(1.2)
org.columns[1].width = Cm(1.8)
org.columns[2].width = Cm(3.2)
org.columns[3].width = Cm(7.5)
org.columns[4].width = Cm(3.0)

headers = ['순번','직위','담당 업무','보유역량 (경력 및 학력 등)','구성 상태']
for i, h in enumerate(headers):
    head_cell(org.cell(0,i), h)
set_repeat_header(org.rows[0])

for ri, row_data in enumerate(org_data):
    r = ri + 1
    cell_para(org.cell(r,0), row_data[0], align='center')
    cell_para(org.cell(r,1), row_data[1], align='center')
    cell_para(org.cell(r,2), row_data[2])
    cell_para(org.cell(r,3), row_data[3])
    cell_para(org.cell(r,4), row_data[4], align='center')


# ════════════════════
# □ 창업아이템 개요
# ════════════════════
section_title(doc, '창업아이템 개요 (요약)')
caption(doc, '＜ 창업아이템 개요 ＞')

ov = doc.add_table(rows=6, cols=4)
set_table_style(ov)
ov.columns[0].width = Cm(3.2)
ov.columns[1].width = Cm(5.5)
ov.columns[2].width = Cm(2.5)
ov.columns[3].width = Cm(6.0)

head_cell(ov.cell(0,0), '명칭')
cell_para(ov.cell(0,1), '오토케 (AutoKe)')
head_cell(ov.cell(0,2), '범주')
cell_para(ov.cell(0,3), '자동화 소프트웨어 (AI 기반)')

# Merge cols 1-3 for rows 1-5
def merge_right(tbl, row, from_col=1, to_col=3):
    c = tbl.cell(row, from_col)
    c.merge(tbl.cell(row, to_col))
    return c

head_cell(ov.cell(1,0), '아이템 개요')
merge_right(ov,1)
cell_para(ov.cell(1,1), '마케팅 콘텐츠를 자동으로 생성하고, SNS에 자동으로 스케줄링·게시하며, 성과 데이터를 분석해 최적화까지 수행하는 마케팅 초보자를 위한 통합 자동화 플랫폼')

head_cell(ov.cell(2,0), '핵심 기능')
merge_right(ov,2)
cell_para(ov.cell(2,1), '콘텐츠 자동 생성 → SNS 스케줄링·게시 → 데이터 기반 성과 분석 → 최적화 (End-to-End 마케팅 자동화)')

head_cell(ov.cell(3,0), '대상 고객')
merge_right(ov,3)
cell_para(ov.cell(3,1), '1인 창업자 약 100만 명, 소상공인 약 600만 명')

head_cell(ov.cell(4,0), '고객 혜택')
merge_right(ov,4)
cell_para(ov.cell(4,1), '시간 절약, 운영 자동화, 마케팅 효율 극대화')

head_cell(ov.cell(5,0), '가격 모델')
merge_right(ov,5)
cell_para(ov.cell(5,1), '월 구독제 (SaaS) + 토큰형 콘텐츠 생성 과금 병행')

# 핵심 기능 및 특징 표
caption(doc, '＜ 핵심 기능 및 특징 ＞')
ft = doc.add_table(rows=5, cols=4)
set_table_style(ft)
ft.columns[0].width = Cm(2.8)
ft.columns[1].width = Cm(3.5)
ft.columns[2].width = Cm(6.5)
ft.columns[3].width = Cm(4.0)

for i, h in enumerate(['구분','핵심 기능','설명','특징']):
    head_cell(ft.cell(0,i), h)
set_repeat_header(ft.rows[0])

ft_rows = [
    ('콘텐츠 생성','AI 기반 텍스트·이미지 자동 생성','제품 정보 입력 시 SNS 규격 콘텐츠 자동 생성. 단시간 내 100개 콘텐츠 생성 가능','플랫폼별 포맷 대응, 이미지+텍스트 혼합'),
    ('자동 스케줄링','SNS 채널별 시간 자동 추천 및 게시','사용자 활동 시간대 분석 → 자동 업로드 실행. 날짜·요일·시간 자동 반복 스케줄링','Instagram, Threads, Twitter(X), Facebook'),
    ('분석·전략 추천','콘텐츠 반응 기반 성과 분석 + 전략 제안','좋아요·조회수 등 KPI 기반 전략 자동 추천. 최적 게시 시간 및 콘텐츠 유형 추천','대시보드 리포트 자동 생성'),
    ('대시보드','관리자용 운영·성과 시각화 UI','콘텐츠별 요일·시간에 따른 성과 실시간 확인. 직관적 UX로 초보 창업자 친화 설계','직관적 UX, 성과 기반 콘텐츠 스케줄링 지원'),
]
for ri, (g, f, d, t) in enumerate(ft_rows):
    r = ri+1
    head_cell(ft.cell(r,0), g, size=9.5)
    cell_para(ft.cell(r,1), f)
    cell_para(ft.cell(r,2), d)
    cell_para(ft.cell(r,3), t)


# ════════════════════
# □ 배경 및 필요성
# ════════════════════
section_title(doc, '1-1. 창업아이템 배경 및 필요성')
caption(doc, '＜ 문제 인식 (Problem) ＞')

pm_data = [
    # (구분, 항목, 내용)
    ('시장 문제','높은 제작 비용','SNS 콘텐츠 1개 제작 시 최소 10시간 이상 소요. 창업 초기 기업은 마케팅 인력이 부족하며 반복 작업으로 본업 집중 어려움'),
    ('시장 문제','도구 파편화','콘텐츠 생성 플랫폼과 마케팅 자동화 도구가 서로 분리되어 있어 통합 솔루션 부재. 자동화 도구에 대한 수요는 높지만 진입장벽이 높음'),
    ('시장 문제','언어 장벽','대부분 외산 위주로 한국어를 지원하지 않아 접근성 및 활용성이 매우 낮음'),
    ('기존 도구 한계','기능 불완전','대부분 콘텐츠 제작만 지원하며 자동 스케줄 기능이 없어 SNS 채널별로 하나하나 직접 설정해야 함'),
    ('기존 도구 한계','효과 검증 불가','마케팅 집행 후 효과 검증이 어려움. 데이터 분석 화면이 전문가 중심으로 창업자가 쓰기 어려운 구조'),
    ('기존 도구 한계','대표자 실증','2회 창업 경험에서 마케팅 실패 반복. 현재 SNS 직접 운영 중으로 콘텐츠 제작에 매주 10시간 이상 소요되는 문제 직접 체감'),
]

prob_tbl = doc.add_table(rows=len(pm_data)+1, cols=3)
set_table_style(prob_tbl)
prob_tbl.columns[0].width = Cm(2.8)
prob_tbl.columns[1].width = Cm(3.5)
prob_tbl.columns[2].width = Cm(11.0)

for i, h in enumerate(['구분','문제 유형','내용']):
    head_cell(prob_tbl.cell(0,i), h)
set_repeat_header(prob_tbl.rows[0])

# 세로병합: 시장 문제 (rows 1-3), 기존 도구 한계 (rows 4-6)
for ri, (g,s,c) in enumerate(pm_data):
    r = ri+1
    if ri == 0:
        merge_v(prob_tbl, 0, 1, 3)
        head_cell(prob_tbl.cell(1,0), '시장\n문제', size=9.5)
    elif ri == 3:
        merge_v(prob_tbl, 0, 4, 6)
        head_cell(prob_tbl.cell(4,0), '기존 도구\n한계', size=9.5)
    subhead_cell(prob_tbl.cell(r,1), s)
    cell_para(prob_tbl.cell(r,2), c)

caption(doc, '＜ 해결방안 및 개발 목적 ＞')

sol_data = [
    ('해결 방안','콘텐츠 자동 생성','입력한 콘텐츠 정보를 통해 단시간 안에 자동으로 100개 콘텐츠 생성'),
    ('해결 방안','자동 스케줄링','생성된 콘텐츠로 4개 SNS(Instagram, Threads, Twitter(X), Facebook)에 자동 스케줄링·게시. 날짜·요일·시간 자동 반복 자동화'),
    ('해결 방안','성과 분석·최적화','콘텐츠의 좋아요·노출수 등을 분석해 스케줄링 최적화 → 마케팅 효율 극대화'),
    ('해결 방안','개발 준비 현황','2023년 AI 모델 실험 → 2024년 소규모 테스트 → 2025년 본격 개발. 1인 창업자·소상공인 5인 인터뷰 완료, 베타 테스트 예정'),
    ('개발 목적','지속 가능한 마케팅','마케팅 전담 인력 없이도 지속 가능한 마케팅 환경 제공'),
    ('개발 목적','비용·시간 절감','콘텐츠 제작 및 게시 업무 간소화로 창업자의 시간·비용 절감'),
    ('개발 목적','다채널 마케팅','다채널 마케팅으로 고객 유입 극대화. 창업자가 마케팅 인력 없이도 유입을 지속시키는 현실적인 솔루션 제공'),
]

sol_tbl = doc.add_table(rows=len(sol_data)+1, cols=3)
set_table_style(sol_tbl)
sol_tbl.columns[0].width = Cm(2.8)
sol_tbl.columns[1].width = Cm(3.5)
sol_tbl.columns[2].width = Cm(11.0)

for i, h in enumerate(['구분','항목','내용']):
    head_cell(sol_tbl.cell(0,i), h)
set_repeat_header(sol_tbl.rows[0])

merge_v(sol_tbl, 0, 1, 4)
head_cell(sol_tbl.cell(1,0), '해결\n방안', size=9.5)
merge_v(sol_tbl, 0, 5, 7)
head_cell(sol_tbl.cell(5,0), '개발\n목적', size=9.5)

for ri, (g,s,c) in enumerate(sol_data):
    r = ri+1
    subhead_cell(sol_tbl.cell(r,1), s)
    cell_para(sol_tbl.cell(r,2), c)


# ════════════════════════
# □ 목표시장 및 사업화 전략
# ════════════════════════
section_title(doc, '목표시장 및 사업화 전략')
caption(doc, '＜ 목표시장 및 사업화 전략 ＞')

biz_data = [
    ('시장 현황','목표 시장','국내 자영업자 약 600만 명 / 1인 창업자 및 크리에이터 시장 급성장. SNS 마케팅은 핵심 채널이지만 실행력 부족으로 성과 저조'),
    ('시장 현황','고객 요구사항','· 콘텐츠를 \'누가 대신\' 만들어 주길 원함\n· 요일·시간 등 분석 및 추천 기능 필요\n· 쉬운 인터페이스와 명확한 성과 피드백 대시보드 필요\n· 적은 리소스로 운영 가능한 마케팅 솔루션 필요'),
    ('비즈니스\n모델','수익 구조','월 구독형 SaaS 기반 + 콘텐츠 생성 토큰 과금 방식 병행 → 콘텐츠는 필요 수만큼 유연 사용'),
    ('비즈니스\n모델','고객 확보 전략','1인 창업자·소상공인 대상 무료 체험 → 유료 전환. 초기 커뮤니티·창업 플랫폼 중심 타겟팅, 실사용 기반 확산'),
    ('성과 목표','가입자','유료 가입자 1,000명 확보'),
    ('성과 목표','매출','월 반복 매출(MRR) 1,000만 원 이상'),
    ('성과 목표','고용','고용 인력 2명 이상'),
    ('핵심\n경쟁력','차별성','콘텐츠 생성부터 게시·분석·최적화까지 마케팅 전체를 자동화. \'오토케(어떻게)\'를 고민하는 사람을 위한 실행 중심 설계'),
    ('핵심\n경쟁력','기술·조직 경쟁력','자체 AI 콘텐츠 생성 실험 경험(2023~), OpenAI 등 다양한 AI를 활용한 실서비스 설계, 한양대 창업대학원 네트워크와 실증 가능'),
]

biz_tbl = doc.add_table(rows=len(biz_data)+1, cols=3)
set_table_style(biz_tbl)
biz_tbl.columns[0].width = Cm(2.8)
biz_tbl.columns[1].width = Cm(3.5)
biz_tbl.columns[2].width = Cm(11.0)

for i, h in enumerate(['구분','항목','내용']):
    head_cell(biz_tbl.cell(0,i), h)
set_repeat_header(biz_tbl.rows[0])

merge_v(biz_tbl, 0, 1, 2);  head_cell(biz_tbl.cell(1,0), '시장\n현황', size=9.5)
merge_v(biz_tbl, 0, 3, 4);  head_cell(biz_tbl.cell(3,0), '비즈니스\n모델', size=9.5)
merge_v(biz_tbl, 0, 5, 7);  head_cell(biz_tbl.cell(5,0), '성과\n목표', size=9.5)
merge_v(biz_tbl, 0, 8, 9);  head_cell(biz_tbl.cell(8,0), '핵심\n경쟁력', size=9.5)

for ri, (g,s,c) in enumerate(biz_data):
    r = ri+1
    subhead_cell(biz_tbl.cell(r,1), s)
    cell_para(biz_tbl.cell(r,2), c)

# 최종 산출물 표
caption(doc, '＜ 최종 산출물 작성 (전체 사업단계) ＞')

del_data = [
    ('단기\n(협약\n기간 내)','AI 텍스트·이미지 콘텐츠 생성 모듈','1','2025.07'),
    ('단기\n(협약\n기간 내)','SNS 자동 스케줄링 모듈','1','2025.07'),
    ('단기\n(협약\n기간 내)','관리자용 대시보드 페이지','1','2025.10'),
    ('단기\n(협약\n기간 내)','마케팅 성과 분석 및 AI 전략 추천 모듈','1','2025.12'),
    ('단기\n(협약\n기간 내)','오토케 웹 기반 플랫폼','1','2025.12'),
    ('단기\n(협약\n기간 내)','AI 기반 마케팅 자동화 기술 특허 출원','1','2025.12'),
    ('중·장기\n최종\n산출물','브랜드 맞춤형 AI 전략 추천 알고리즘 모듈','1','2026.06'),
    ('중·장기\n최종\n산출물','외부 플랫폼 연동 API 모듈','1','2026.12'),
    ('중·장기\n최종\n산출물','온라인 광고 자동 집행 모듈\n(구글 애즈, 네이버광고, 카카오모먼트 연동 포함)','1','2026.12'),
    ('중·장기\n최종\n산출물','오토케 앱 (iOS, AOS 웹 플랫폼 기능 앱과 연동)','1','2026.12'),
]

del_tbl = doc.add_table(rows=len(del_data)+1, cols=4)
set_table_style(del_tbl)
del_tbl.columns[0].width = Cm(2.8)
del_tbl.columns[1].width = Cm(9.5)
del_tbl.columns[2].width = Cm(1.5)
del_tbl.columns[3].width = Cm(2.5)

for i, h in enumerate(['구분','최종 산출물','수량','완료 일정']):
    head_cell(del_tbl.cell(0,i), h)
set_repeat_header(del_tbl.rows[0])

merge_v(del_tbl, 0, 1, 6);  head_cell(del_tbl.cell(1,0), '단기\n(협약\n기간 내)', size=9.5)
merge_v(del_tbl, 0, 7, 10); head_cell(del_tbl.cell(7,0), '중·장기\n최종\n산출물', size=9.5)

for ri, (g,s,c,d) in enumerate(del_data):
    r = ri+1
    cell_para(del_tbl.cell(r,1), s)
    cell_para(del_tbl.cell(r,2), c, align='center')
    cell_para(del_tbl.cell(r,3), d, align='center')

# 마지막 여백
doc.add_paragraph()
ep = doc.add_paragraph('— 오토케(AutoKe) 창업사업화 지원사업 사업계획서 —')
ep.alignment = WD_ALIGN_PARAGRAPH.CENTER
ep.paragraph_format.space_before = Pt(12)
for run in ep.runs:
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x77,0x77,0x77)

# ── 저장 ──
out_path = '/Users/kim/.openclaw/workspace/enttix/오토케_사업계획서.docx'
doc.save(out_path)
print(f'✅ 저장 완료: {out_path}')
